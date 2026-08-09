import io


# ── Test fixtures ────────────────────────────────────────────────────

def _valid_docx_bytes() -> bytes:
    from docx import Document
    doc = Document()
    doc.add_paragraph("Perennia offers AI consulting and custom automation.")
    doc.add_paragraph("Our office hours are 9am to 5pm, Sunday through Thursday.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _valid_pdf_bytes() -> bytes:
    # Minimal single-page PDF with one line of text, hand-built (no
    # dependency needed to construct it — pypdf only needs to be able
    # to read it back, which this well-formed minimal PDF supports).
    content = b"BT /F1 24 Tf 100 700 Td (Perennia PDF test content) Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 612 792] /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(content), content),
    ]
    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(out.tell())
        out.write(f"{i} 0 obj\n".encode() + obj + b"\nendobj\n")
    xref_offset = out.tell()
    out.write(f"xref\n0 {len(objects) + 1}\n".encode())
    out.write(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        out.write(f"{off:010d} 00000 n \n".encode())
    out.write(f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n".encode())
    out.write(f"startxref\n{xref_offset}\n%%EOF".encode())
    return out.getvalue()


HTML_BYTES = b"<html><head><title>T</title><style>.x{color:red}</style></head>" \
             b"<body><h1>Perennia Pricing</h1><p>Contact us for a custom quote.</p>" \
             b"<script>alert(1)</script></body></html>"

MD_BYTES = b"# Perennia FAQ\n\nWe support English and Arabic.\n"


# ── Extraction unit tests ──────────────────────────────────────────

def test_extract_markdown():
    from app.knowledge_extract import extract_from_upload
    text, truncated, ct = extract_from_upload(MD_BYTES, "faq.md", max_chars=10000)
    assert "Perennia FAQ" in text
    assert ct == "markdown"
    assert truncated is False


def test_extract_plain_text():
    from app.knowledge_extract import extract_from_upload
    text, truncated, ct = extract_from_upload(b"Just plain text content.", "notes.txt", max_chars=10000)
    assert text == "Just plain text content."
    assert ct == "text"


def test_extract_html_strips_scripts_and_styles():
    from app.knowledge_extract import extract_from_upload
    text, truncated, ct = extract_from_upload(HTML_BYTES, "page.html", max_chars=10000)
    assert "Perennia Pricing" in text
    assert "Contact us for a custom quote" in text
    assert "alert(1)" not in text
    assert "color:red" not in text
    assert ct == "html"


def test_extract_docx():
    from app.knowledge_extract import extract_from_upload
    text, truncated, ct = extract_from_upload(_valid_docx_bytes(), "info.docx", max_chars=10000)
    assert "AI consulting" in text
    assert ct == "docx"


def test_extract_pdf():
    from app.knowledge_extract import extract_from_upload
    text, truncated, ct = extract_from_upload(_valid_pdf_bytes(), "info.pdf", max_chars=10000)
    assert "Perennia PDF test content" in text
    assert ct == "pdf"


def test_extract_rejects_unsupported_extension():
    from app.knowledge_extract import ExtractionError, extract_from_upload
    import pytest
    with pytest.raises(ExtractionError):
        extract_from_upload(b"fake exe content", "malware.exe", max_chars=10000)


def test_extract_rejects_mismatched_content():
    """A binary file renamed to .txt must be rejected — never trust
    the extension over the actual bytes."""
    from app.knowledge_extract import ExtractionError, extract_from_upload
    import pytest
    fake_binary = bytes([0, 1, 2, 3, 255, 254, 253] * 20)
    with pytest.raises(ExtractionError):
        extract_from_upload(fake_binary, "renamed.txt", max_chars=10000)


def test_extract_sniffs_actual_content_over_claimed_extension():
    """Content type is determined by sniffing bytes, not trusting the
    filename — a PDF's real content is correctly extracted as a PDF
    even if it arrives with a misleading .docx extension, mirroring
    the same "never trust the client-supplied extension" principle
    used for image uploads."""
    from app.knowledge_extract import extract_from_upload
    text, truncated, ct = extract_from_upload(_valid_pdf_bytes(), "fake.docx", max_chars=10000)
    assert ct == "pdf"
    assert "Perennia PDF test content" in text


def test_extract_docx_sniffed_correctly_even_with_txt_extension():
    """Same principle, the other direction: a real docx's content is
    correctly detected and extracted even if renamed to .txt — sniffing
    always wins over the claimed extension."""
    from app.knowledge_extract import extract_from_upload
    text, truncated, ct = extract_from_upload(_valid_docx_bytes(), "fake.txt", max_chars=10000)
    assert ct == "docx"
    assert "AI consulting" in text


def test_extract_truncates_long_content():
    from app.knowledge_extract import extract_from_upload
    long_text = ("word " * 5000).encode()
    text, truncated, ct = extract_from_upload(long_text, "long.txt", max_chars=100)
    assert len(text) == 100
    assert truncated is True


# ── SSRF guard tests (no real network needed) ──────────────────────

def test_url_validation_rejects_non_http_scheme():
    from app.knowledge_extract import ExtractionError, extract_from_url
    import pytest
    with pytest.raises(ExtractionError):
        extract_from_url("ftp://example.com/file", max_chars=1000)


def test_url_validation_rejects_localhost():
    from app.knowledge_extract import ExtractionError, extract_from_url
    import pytest
    with pytest.raises(ExtractionError):
        extract_from_url("http://localhost:8001/admin", max_chars=1000)


def test_url_validation_rejects_loopback_ip():
    from app.knowledge_extract import ExtractionError, extract_from_url
    import pytest
    with pytest.raises(ExtractionError):
        extract_from_url("http://127.0.0.1/", max_chars=1000)


def test_url_validation_rejects_private_ip():
    from app.knowledge_extract import ExtractionError, extract_from_url
    import pytest
    with pytest.raises(ExtractionError):
        extract_from_url("http://192.168.1.1/", max_chars=1000)


def test_url_validation_rejects_link_local_metadata_ip():
    """Blocks the cloud-provider metadata endpoint address — a classic
    SSRF target for credential theft."""
    from app.knowledge_extract import ExtractionError, extract_from_url
    import pytest
    with pytest.raises(ExtractionError):
        extract_from_url("http://169.254.169.254/latest/meta-data/", max_chars=1000)


# ── Service + admin API tests ────────────────────────────────────────

def test_upload_source_requires_auth(client):
    resp = client.post("/admin/api/knowledge/upload", files={"file": ("a.txt", io.BytesIO(b"hi"), "text/plain")})
    assert resp.status_code == 401


def test_upload_and_list_source(logged_in_client):
    resp = logged_in_client.post("/admin/api/knowledge/upload",
                                  files={"file": ("faq.md", io.BytesIO(MD_BYTES), "text/markdown")})
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["ok"] is True
    assert body["content_type"] == "markdown"
    assert body["kind"] == "file"

    listed = logged_in_client.get("/admin/api/knowledge").json()
    assert any(s["id"] == body["id"] for s in listed)
    # list view omits full text
    assert "text" not in listed[0]


def test_get_source_includes_text_preview(logged_in_client):
    created = logged_in_client.post("/admin/api/knowledge/upload",
                                     files={"file": ("faq2.md", io.BytesIO(MD_BYTES), "text/markdown")}).json()
    detail = logged_in_client.get(f"/admin/api/knowledge/{created['id']}").json()
    assert "Perennia FAQ" in detail["text"]


def test_upload_bad_file_creates_failed_source_not_error(logged_in_client):
    """An extraction failure still creates a source record (marked
    not-ok, with an error message) rather than a 500 — the admin sees
    what went wrong in the list instead of a bare failed request."""
    resp = logged_in_client.post("/admin/api/knowledge/upload",
                                  files={"file": ("bad.xyz", io.BytesIO(b"whatever"), "application/octet-stream")})
    assert resp.status_code == 200
    body = resp.json()
    assert body["ok"] is False
    assert body["error_message"]


def test_delete_source(logged_in_client):
    created = logged_in_client.post("/admin/api/knowledge/upload",
                                     files={"file": ("del.txt", io.BytesIO(b"delete me"), "text/plain")}).json()
    resp = logged_in_client.delete(f"/admin/api/knowledge/{created['id']}")
    assert resp.status_code == 200
    assert logged_in_client.get(f"/admin/api/knowledge/{created['id']}").status_code == 404


def test_toggle_active(logged_in_client):
    created = logged_in_client.post("/admin/api/knowledge/upload",
                                     files={"file": ("toggle.txt", io.BytesIO(b"content"), "text/plain")}).json()
    resp = logged_in_client.patch(f"/admin/api/knowledge/{created['id']}", json={"is_active": False})
    assert resp.status_code == 200
    assert resp.json()["is_active"] is False


def test_add_url_rejects_private_address(logged_in_client):
    resp = logged_in_client.post("/admin/api/knowledge/url", json={"url": "http://127.0.0.1/secret"})
    assert resp.status_code == 200  # source is created but marked not-ok, matching upload's error pattern
    assert resp.json()["ok"] is False


def test_refresh_only_works_for_url_sources(logged_in_client):
    created = logged_in_client.post("/admin/api/knowledge/upload",
                                     files={"file": ("f.txt", io.BytesIO(b"x"), "text/plain")}).json()
    resp = logged_in_client.post(f"/admin/api/knowledge/{created['id']}/refresh")
    assert resp.status_code == 400


def test_capacity_limit_enforced(logged_in_client):
    resp = logged_in_client.put("/admin/api/settings/knowledge", json={"knowledge.max_total_sources": 1})
    assert resp.status_code == 200
    try:
        logged_in_client.post("/admin/api/knowledge/upload",
                               files={"file": ("cap1.txt", io.BytesIO(b"one"), "text/plain")})
        # This is the 2nd source overall in a fresh-enough test DB context;
        # capacity is checked against the TOTAL existing count, so we
        # only assert the eventual over-limit case raises 400 somewhere
        # in a tight loop rather than asserting an exact count (other
        # tests share this DB across the session).
        over_resp = None
        for i in range(5):
            r = logged_in_client.post("/admin/api/knowledge/upload",
                                       files={"file": (f"cap{i}.txt", io.BytesIO(b"x"), "text/plain")})
            if r.status_code == 400:
                over_resp = r
                break
        assert over_resp is not None
        assert "full" in over_resp.json()["detail"].lower()
    finally:
        logged_in_client.put("/admin/api/settings/knowledge", json={"knowledge.max_total_sources": 20})


# ── Chat integration ────────────────────────────────────────────────

def test_knowledge_included_in_chat_system_prompt(logged_in_client, client, monkeypatch):
    logged_in_client.post("/admin/api/knowledge/upload",
                           files={"file": ("kbtest.md", io.BytesIO(b"# Special Fact\nThe secret code is BANANA42."),
                                           "text/markdown")})

    captured = {}

    def fake_generate_reply(**kwargs):
        captured.update(kwargs)
        return "canned reply"

    monkeypatch.setattr("app.chat_service.llm_client.generate_reply", fake_generate_reply)
    logged_in_client.put("/admin/api/settings/chat", json={"chat.llm_provider": "anthropic", "chat.llm_api_key": "sk-fake"})
    try:
        resp = client.post("/api/chat", json={"message": "tell me something", "lang": "en", "history": []})
        assert resp.status_code == 200
        assert "BANANA42" in captured["system_prompt"]
        assert "DOCUMENT START: kbtest.md" in captured["system_prompt"]
    finally:
        logged_in_client.put("/admin/api/settings/chat", json={"chat.llm_provider": "none"})


def test_knowledge_disabled_setting_excludes_from_prompt(logged_in_client, client, monkeypatch):
    logged_in_client.post("/admin/api/knowledge/upload",
                           files={"file": ("kbtest2.md", io.BytesIO(b"UNIQUEMARKER998877"), "text/markdown")})
    logged_in_client.put("/admin/api/settings/knowledge", json={"knowledge.enabled": False})

    captured = {}

    def fake_generate_reply(**kwargs):
        captured.update(kwargs)
        return "canned reply"

    monkeypatch.setattr("app.chat_service.llm_client.generate_reply", fake_generate_reply)
    logged_in_client.put("/admin/api/settings/chat", json={"chat.llm_provider": "anthropic", "chat.llm_api_key": "sk-fake"})
    try:
        client.post("/api/chat", json={"message": "hi", "lang": "en", "history": []})
        assert "UNIQUEMARKER998877" not in captured.get("system_prompt", "")
    finally:
        logged_in_client.put("/admin/api/settings/chat", json={"chat.llm_provider": "none"})
        logged_in_client.put("/admin/api/settings/knowledge", json={"knowledge.enabled": True})


def test_inactive_source_excluded_from_prompt(logged_in_client, client, monkeypatch):
    created = logged_in_client.post(
        "/admin/api/knowledge/upload",
        files={"file": ("kbtest3.md", io.BytesIO(b"HIDDENMARKER554433"), "text/markdown")},
    ).json()
    logged_in_client.patch(f"/admin/api/knowledge/{created['id']}", json={"is_active": False})

    captured = {}
    monkeypatch.setattr("app.chat_service.llm_client.generate_reply",
                         lambda **kw: captured.update(kw) or "reply")
    logged_in_client.put("/admin/api/settings/chat", json={"chat.llm_provider": "anthropic", "chat.llm_api_key": "sk-fake"})
    try:
        client.post("/api/chat", json={"message": "hi", "lang": "en", "history": []})
        assert "HIDDENMARKER554433" not in captured.get("system_prompt", "")
    finally:
        logged_in_client.put("/admin/api/settings/chat", json={"chat.llm_provider": "none"})


# ── Settings validation ────────────────────────────────────────────

def test_knowledge_settings_validation(logged_in_client):
    resp = logged_in_client.put("/admin/api/settings/knowledge", json={"knowledge.max_total_sources": 0})
    assert resp.status_code == 400
    resp2 = logged_in_client.put("/admin/api/settings/knowledge", json={"knowledge.max_lines_in_prompt": 1})
    assert resp2.status_code == 400
